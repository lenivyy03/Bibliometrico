from __future__ import annotations

import io
import zipfile
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd

from gui_utils import (FONT_FAMILY, FONT_SM, FONT_MD, FONT_LG, PAD_HEADER, PAD_CARD,
                       ColorButton, bind_mousewheel)
from compat_imports import load_project_module

# ── Backend modules (para cálculo bajo demanda cuando el getter no existe) ──
_conteo_mod = load_project_module("conteo")
_filtrado_mod = load_project_module("filtrado")
_impacto_mod = load_project_module("impacto")
_tops_mod = load_project_module("tops")

BG = "#f5f5f5"
CARD = "#ffffff"
TEXT = "#1f2937"
MUTED = "#6b7280"
ACCENT = "#7c3aed"
BORDER = "#e5e7eb"
SUCCESS = "#16a34a"

_ITEMS = [
    ("metricas",      "Métricas generales",        "Productividad"),
    ("autoria_est",   "Estadísticas de autoría",   "Autoría"),
    ("top_autores",   "Top 10 autores",            "Autoría"),
    ("paises",        "Lista de países",           "Geografía"),
    ("paises_top10",  "Top 10 países",             "Geografía"),
    ("univ_lista",    "Ranking universidades",     "Geografía"),
    ("univ_top",      "Top 10 universidades",      "Geografía"),
    ("citas",         "Promedio anual de citas",   "Impacto"),
    ("top_trabajos",  "Top 10 trabajos",           "Impacto"),
    ("apa",           "Referencias APA",           "General"),
]

_LABELS = {k: l for k, l, _ in _ITEMS}


class VistaExportar(tk.Frame):
    def __init__(self, parent: tk.Widget, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # Header
        head = tk.Frame(self, bg=BG)
        head.grid(row=0, column=0, sticky="ew", padx=24, pady=PAD_HEADER)
        head.grid_columnconfigure(0, weight=1)
        tk.Label(head, text="Exportar", bg=BG, fg=TEXT,
                 font=(FONT_FAMILY, FONT_LG, "bold")).grid(row=0, column=0, sticky="w")
        tk.Label(
            head,
            text="Selecciona las secciones a incluir. Puedes marcarlas desde cada sección o directamente aquí.",
            bg=BG, fg=MUTED, font=(FONT_FAMILY, FONT_SM),
        ).grid(row=1, column=0, sticky="w")

        # Body: left list + right preview
        body = tk.Frame(self, bg=BG)
        body.grid(row=1, column=0, sticky="nsew", padx=24, pady=(0, 10))
        body.grid_columnconfigure(0, weight=2)
        body.grid_columnconfigure(1, weight=3)
        body.grid_rowconfigure(0, weight=1)

        # ── Left panel — lista con scroll ────────────────────────────────────
        left_outer = tk.Frame(body, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        left_outer.grid(row=0, column=0, sticky="nsew", padx=(0, 12))
        left_outer.grid_columnconfigure(0, weight=1)
        left_outer.grid_rowconfigure(1, weight=1)

        tk.Label(left_outer, text="Secciones disponibles", bg=CARD, fg=TEXT,
                 font=(FONT_FAMILY, FONT_MD, "bold")).grid(row=0, column=0, columnspan=2, sticky="w", padx=16, pady=(10, 4))

        left_canvas = tk.Canvas(left_outer, bg=CARD, highlightthickness=0)
        left_canvas.grid(row=1, column=0, sticky="nsew")
        left_sb = ttk.Scrollbar(left_outer, orient="vertical", command=left_canvas.yview)
        left_sb.grid(row=1, column=1, sticky="ns")
        left_canvas.configure(yscrollcommand=left_sb.set)

        left_inner = tk.Frame(left_canvas, bg=CARD)
        left_win = left_canvas.create_window((0, 0), window=left_inner, anchor="nw")
        left_inner.bind("<Configure>", lambda _e: left_canvas.configure(scrollregion=left_canvas.bbox("all")))
        left_canvas.bind("<Configure>", lambda e: left_canvas.itemconfigure(left_win, width=e.width))
        bind_mousewheel(left_canvas, left_canvas)
        bind_mousewheel(left_inner, left_canvas)

        self._item_btns: dict[str, ColorButton] = {}
        for key, label, cat in _ITEMS:
            fila = tk.Frame(left_inner, bg=CARD)
            fila.pack(fill="x", padx=12, pady=PAD_CARD)

            btn = ColorButton(
                fila, text="○",
                bg="#f3f4f6", fg=MUTED,
                relief="flat", cursor="hand2",
                font=(FONT_FAMILY, FONT_MD + 1, "bold"),
                padx=6, pady=2, width=2,
            )
            btn.configure(command=lambda k=key: self._toggle(k))
            btn.pack(side="left", padx=(0, 10))
            self._item_btns[key] = btn

            info = tk.Frame(fila, bg=CARD)
            info.pack(side="left", fill="x", expand=True)
            tk.Label(info, text=label, bg=CARD, fg=TEXT,
                     font=(FONT_FAMILY, FONT_MD, "bold"), anchor="w").pack(anchor="w")
            tk.Label(info, text=cat, bg=CARD, fg=MUTED,
                     font=(FONT_FAMILY, FONT_SM), anchor="w").pack(anchor="w")

            bind_mousewheel(fila, left_canvas)
            bind_mousewheel(info, left_canvas)

        # ── Right panel — scrollable preview ────────────────────────────────
        right_outer = tk.Frame(body, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        right_outer.grid(row=0, column=1, sticky="nsew")
        right_outer.grid_columnconfigure(0, weight=1)
        right_outer.grid_rowconfigure(1, weight=1)

        tk.Label(right_outer, text="Vista previa del reporte", bg=CARD, fg=TEXT,
                 font=(FONT_FAMILY, FONT_MD, "bold")).grid(row=0, column=0, sticky="w", padx=16, pady=(10, 6))

        right_canvas = tk.Canvas(right_outer, bg=CARD, highlightthickness=0)
        right_canvas.grid(row=1, column=0, sticky="nsew")
        right_sb = ttk.Scrollbar(right_outer, orient="vertical", command=right_canvas.yview)
        right_sb.grid(row=1, column=1, sticky="ns")
        right_canvas.configure(yscrollcommand=right_sb.set)

        self._preview_frame = tk.Frame(right_canvas, bg=CARD)
        right_win = right_canvas.create_window((0, 0), window=self._preview_frame, anchor="nw")
        self._preview_frame.bind("<Configure>", lambda _e: right_canvas.configure(scrollregion=right_canvas.bbox("all")))
        right_canvas.bind("<Configure>", lambda e: (
            right_canvas.itemconfigure(right_win, width=e.width),
            self._update_wraplength(e.width - 40),
        ))
        bind_mousewheel(right_canvas, right_canvas)
        bind_mousewheel(self._preview_frame, right_canvas)
        self._right_canvas = right_canvas
        self._preview_frame.grid_columnconfigure(0, weight=1)
        self._preview_wraplength = 340

        # Bottom — format buttons
        bottom = tk.Frame(self, bg=BG)
        bottom.grid(row=2, column=0, sticky="ew", padx=24, pady=(0, 24))
        self.status_label = tk.Label(bottom, text="", bg=BG, fg=MUTED, font=(FONT_FAMILY, FONT_MD))
        self.status_label.pack(anchor="w", pady=(0, 8))
        acciones = tk.Frame(bottom, bg=BG)
        acciones.pack(anchor="w")
        ColorButton(
            acciones, text="Exportar como Excel",
            command=lambda: self._exportar("excel"),
            bg=ACCENT, fg="white", relief="flat", cursor="hand2",
            font=(FONT_FAMILY, FONT_MD, "bold"), padx=16, pady=8,
        ).pack(side="left")
        ColorButton(
            acciones, text="Exportar como Word",
            command=lambda: self._exportar("word"),
            bg="#ede9fe", fg=ACCENT, relief="flat", cursor="hand2",
            font=(FONT_FAMILY, FONT_MD, "bold"), padx=16, pady=8,
        ).pack(side="left", padx=(10, 0))
        ColorButton(
            acciones, text="Exportar como CSV (.zip)",
            command=lambda: self._exportar("csv"),
            bg="#f3f4f6", fg=TEXT, relief="flat", cursor="hand2",
            font=(FONT_FAMILY, FONT_MD, "bold"), padx=16, pady=8,
        ).pack(side="left", padx=(10, 0))

        # Register callbacks so external toggles update our buttons too
        for key, _, _ in _ITEMS:
            self.app.registrar_toggle_cb(key, lambda sel, k=key: self._on_toggle_cb(k, sel))

        self._actualizar_preview()

    def _update_wraplength(self, width: int) -> None:
        self._preview_wraplength = max(200, width)
        self._actualizar_preview()

    # ── Toggle logic ──────────────────────────────────────────────────────────

    def _toggle(self, key: str) -> None:
        self.app.toggle_exportacion(key)

    def _on_toggle_cb(self, key: str, seleccionado: bool) -> None:
        btn = self._item_btns.get(key)
        if btn:
            if seleccionado:
                btn.configure(text="●", bg=SUCCESS, fg="white")
            else:
                btn.configure(text="○", bg="#f3f4f6", fg=MUTED)
        self._actualizar_preview()

    # ── Preview ───────────────────────────────────────────────────────────────

    def _actualizar_preview(self) -> None:
        for w in self._preview_frame.winfo_children():
            w.destroy()

        seleccionados = [(k, l) for k, l, _ in _ITEMS if k in self.app.export_selections]

        if not seleccionados:
            tk.Label(
                self._preview_frame,
                text="Ninguna sección seleccionada.\n\nUsa los botones  Exportar ✓  en cada sección,\no marca directamente aquí con los círculos.",
                bg=CARD, fg=MUTED, font=(FONT_FAMILY, FONT_MD), justify="left",
                wraplength=self._preview_wraplength, anchor="nw",
            ).grid(row=0, column=0, sticky="nw")
            self.status_label.configure(text="")
            return

        n = len(seleccionados)
        tk.Label(
            self._preview_frame,
            text=f"{n} sección{'es' if n > 1 else ''} seleccionada{'s' if n > 1 else ''}:",
            bg=CARD, fg=TEXT, font=(FONT_FAMILY, FONT_MD, "bold"),
        ).grid(row=0, column=0, sticky="w", pady=(0, 10))

        for i, (key, label) in enumerate(seleccionados):
            desc = self._describir(key)
            fila = tk.Frame(self._preview_frame, bg=CARD)
            fila.grid(row=i + 1, column=0, sticky="ew", pady=3)
            tk.Label(fila, text="✓", bg=CARD, fg=SUCCESS,
                     font=(FONT_FAMILY, FONT_MD, "bold"), width=2).pack(side="left")
            tk.Label(fila, text=f"{label}  —  {desc}", bg=CARD, fg=TEXT,
                     font=(FONT_FAMILY, FONT_MD), wraplength=self._preview_wraplength - 20,
                     justify="left").pack(side="left", fill="x")

        self.status_label.configure(text=f"{n} sección(es) lista(s) para exportar.")

    def _describir(self, key: str) -> str:
        if self.app.df is None:
            return "sin datos cargados"
        if key in ("top_autores", "apa"):
            return f"{len(self.app.df)} artículos disponibles"

        # Intentar getter registrado primero, luego fallback
        getter = self.app.export_getters.get(key)
        data = None
        if getter is not None:
            try:
                data = getter()
            except Exception:
                data = None

        if data is None:
            # Describir sin computar todo: si hay df cargado, los datos se pueden generar
            return "datos disponibles al exportar"

        if isinstance(data, dict):
            return f"{len(data)} métricas"
        if isinstance(data, pd.DataFrame):
            return f"{len(data)} registros"
        if isinstance(data, list):
            return f"{len(data)} elementos"
        return "datos disponibles"

    # ── Data getters ──────────────────────────────────────────────────────────

    def _get_df_for_key(self, key: str):
        """Obtiene un DataFrame listo para exportar.

        Primero intenta usar el getter registrado por la vista correspondiente
        (datos ya calculados y en memoria). Si no existe (el usuario nunca
        visitó esa sección), calcula los datos bajo demanda usando las mismas
        funciones del backend que usan las vistas individuales.
        """
        # Secciones que siempre se computan directamente desde el df global
        if key == "top_autores":
            return self._compute_top_autores()
        if key == "apa":
            return self._compute_apa()

        # Intentar getter registrado por la vista
        getter = self.app.export_getters.get(key)
        data = None
        if getter is not None:
            try:
                data = getter()
            except Exception:
                data = None

        # ── Fallback: calcular bajo demanda si no hay getter o devolvió None ─
        if data is None:
            data = self._compute_fallback(key)

        if data is None:
            return None

        # ── Convertir al formato DataFrame uniforme ──────────────────────────
        if key == "metricas" and isinstance(data, dict):
            return pd.DataFrame(list(data.items()), columns=["Métrica", "Valor"])
        if key == "paises" and isinstance(data, list):
            return pd.DataFrame(data, columns=["País", "Artículos"])
        if key == "paises_top10" and isinstance(data, list):
            return pd.DataFrame(data[:10], columns=["País", "Artículos"])
        if key == "top_trabajos" and isinstance(data, list):
            return pd.DataFrame(data, columns=["Referencia_APA", "Citas", "Año", "DOI"])
        if isinstance(data, pd.DataFrame):
            if key == "univ_top":
                return data.head(10)
            return data
        return None

    def _compute_fallback(self, key: str):
        """Calcula datos bajo demanda cuando el usuario no ha visitado la sección.

        Usa exactamente las mismas funciones del backend que usan las vistas
        individuales, garantizando consistencia en los resultados.
        """
        df = self.app.df
        if df is None or df.empty:
            return None
        try:
            if key == "metricas":
                return _conteo_mod.metricas_prod(df)

            if key == "autoria_est":
                return _filtrado_mod.tabla_comparativa_autoria(df)

            if key in ("paises", "paises_top10"):
                paises = _tops_mod.extraer_paises(df)
                freqs = _tops_mod.obtener_frecuencias_paises(paises)
                return [(pais, int(conteo)) for pais, conteo in freqs.items()]

            if key == "univ_lista":
                return _filtrado_mod.obtener_ranking_universidades(df)

            if key == "univ_top":
                return _filtrado_mod.obtener_top_10_universidades_citadas(df)

            if key == "citas":
                tabla = _impacto_mod.calcular_promedio_citas_anual(df)
                return _impacto_mod.ordenar_por_promedio_citas(tabla)

            if key == "top_trabajos":
                return list(_tops_mod.top_10_trabajos(df))

        except Exception:
            return None
        return None

    def _compute_top_autores(self):
        """HU #53 y #36: formato expandido — una fila por título.
        Col A: Autor | Col B: Título del artículo | Col C: Total artículos del autor.
        """
        df = self.app.df
        if df is None or df.empty or "Authors" not in df.columns or "Title" not in df.columns:
            return None
        datos = []
        for _, fila in df.dropna(subset=["Authors", "Title"]).iterrows():
            titulo = str(fila["Title"]).strip()
            for autor in str(fila["Authors"]).split(";"):
                autor = autor.strip()
                if autor:
                    datos.append({"Autor": autor, "Título del artículo": titulo})
        if not datos:
            return None
        df_plano = pd.DataFrame(datos)
        conteos = df_plano.groupby("Autor")["Título del artículo"].nunique()
        top_10 = conteos.nlargest(10).index
        df_top = df_plano[df_plano["Autor"].isin(top_10)].copy()
        df_top["Total artículos"] = df_top["Autor"].map(conteos)
        df_top = df_top.sort_values(
            ["Total artículos", "Autor"], ascending=[False, True]
        ).reset_index(drop=True)
        return df_top[["Autor", "Título del artículo", "Total artículos"]]

    def _compute_apa(self):
        df = self.app.df
        if df is None or not {"Authors", "Title", "Year"}.issubset(df.columns):
            return None
        df_apa = df.copy()
        df_apa["Referencia_APA"] = (
            df_apa["Authors"].fillna("Sin autor") + ". (" +
            df_apa["Year"].astype(str) + "). " +
            df_apa["Title"].fillna("Sin título") + "."
        )
        if "DOI" in df_apa.columns:
            df_apa["Referencia_APA"] += df_apa["DOI"].apply(
                lambda doi: f" https://doi.org/{doi}" if pd.notna(doi) else ""
            )
        return df_apa[["Referencia_APA"]]

    # ── Export ────────────────────────────────────────────────────────────────

    def _exportar(self, formato: str) -> None:
        seleccionados = [k for k, _, _ in _ITEMS if k in self.app.export_selections]
        if not seleccionados:
            messagebox.showwarning("Sin selección", "Selecciona al menos una sección para exportar.")
            return

        if formato == "excel":
            ext, tipos, init = ".xlsx", [("Excel", "*.xlsx")], "reporte_bibliometrico"
        elif formato == "word":
            ext, tipos, init = ".docx", [("Word", "*.docx")], "reporte_bibliometrico"
        else:
            ext, tipos, init = ".zip", [("ZIP con CSVs", "*.zip")], "reporte_bibliometrico"

        ruta = filedialog.asksaveasfilename(defaultextension=ext, filetypes=tipos, initialfile=init)
        if not ruta:
            return

        try:
            if formato == "excel":
                self._exportar_excel(seleccionados, ruta)
            elif formato == "word":
                self._exportar_word(seleccionados, ruta)
            else:
                self._exportar_csv_zip(seleccionados, ruta)
            messagebox.showinfo("Exportación exitosa", f"Reporte guardado:\n{ruta}")
        except Exception as exc:
            messagebox.showerror("Error al exportar", str(exc))

    def _exportar_excel(self, claves: list, ruta: str) -> None:
        with pd.ExcelWriter(ruta, engine="openpyxl") as writer:
            for key in claves:
                df = self._get_df_for_key(key)
                if df is None or df.empty:
                    continue
                sheet = _LABELS.get(key, key)[:31]
                df.to_excel(writer, sheet_name=sheet, index=False)

    def _exportar_word(self, claves: list, ruta: str) -> None:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "La librería python-docx no está instalada.\n"
                "Instálala con:  pip install python-docx"
            )
        doc = Document()
        doc.add_heading("Reporte Bibliométrico", level=0)
        for key in claves:
            df = self._get_df_for_key(key)
            if df is None or df.empty:
                continue
            doc.add_heading(_LABELS.get(key, key), level=1)
            tabla = doc.add_table(rows=1, cols=df.shape[1])
            tabla.style = "Table Grid"
            for i, col in enumerate(df.columns):
                tabla.rows[0].cells[i].text = str(col)
            for _, row in df.iterrows():
                cells = tabla.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            doc.add_paragraph()
        doc.save(ruta)

    def _exportar_csv_zip(self, claves: list, ruta: str) -> None:
        BOM = b'\xef\xbb\xbf'  # UTF-8 BOM — necesario para que Excel abra sin errores
        with zipfile.ZipFile(ruta, "w", zipfile.ZIP_DEFLATED) as zf:
            for key in claves:
                df = self._get_df_for_key(key)
                if df is None or df.empty:
                    continue
                nombre = _LABELS.get(key, key).replace(" ", "_") + ".csv"
                csv_bytes = BOM + df.to_csv(index=False).encode("utf-8")
                zf.writestr(nombre, csv_bytes)

    def on_show(self) -> None:
        for key, _, _ in _ITEMS:
            sel = key in self.app.export_selections
            btn = self._item_btns.get(key)
            if btn:
                if sel:
                    btn.configure(text="●", bg=SUCCESS, fg="white")
                else:
                    btn.configure(text="○", bg="#f3f4f6", fg=MUTED)
        self._actualizar_preview()
