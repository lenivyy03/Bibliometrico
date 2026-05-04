import pandas as pd
from docx import Document


# --- Exportar Universidades (Historia 24 y 40) ---

def exportar_lista_universidades(datos_resultados: pd.DataFrame, nombre_archivo: str = "resultado_bibliometria",
                                 formato: str = "csv") -> str | None:
    datos_preparados = datos_resultados.copy()

    for columnas in datos_preparados.columns:
        if len(datos_preparados) > 0 and isinstance(datos_preparados[columnas].iloc[0], list):
            datos_preparados[columnas] = datos_preparados[columnas].apply(
                lambda lista_valores: " | ".join(map(str, lista_valores))
            )

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        datos_preparados.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        datos_preparados.to_excel(archivo_salida, index=False)
    else:
        print("Formato no soportado. Elige 'csv' o 'xlsx'.")
        return None

    print(f"Archivo exportado exitosamente como: {archivo_salida}")
    return archivo_salida


def exportar_top_10_universidades(df_top_10_universidades: pd.DataFrame, nombre_archivo: str = "top_10_universidades",
                                  formato: str = "csv") -> str | None:
    datos_exportar = df_top_10_universidades.copy()

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        datos_exportar.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        datos_exportar.to_excel(archivo_salida, index=False)
    else:
        print("Formato no soportado. Elige 'csv' o 'xlsx'.")
        return None

    print(f"Archivo del top 10 exportado exitosamente como: {archivo_salida}")
    return archivo_salida


# --- Exportar Países (Historia 43) ---

def exportar_top_10_paises_excel(df_top_10_paises: pd.DataFrame, nombre_archivo: str = "ranking_top_10_paises") -> str:
    archivo_salida = f"{nombre_archivo}.xlsx"
    df_top_10_paises.to_excel(archivo_salida, index=False)
    print(f"Ranking de países exportado exitosamente a Excel como: {archivo_salida}")
    return archivo_salida


# --- Exportar Conteos, Autoría e Impacto (Historias 39, 40, 41) ---

def exportar_resumen_conteos(resumen: dict, nombre_archivo: str = "resumen_conteos",
                             formato: str = "csv") -> str | None:
    if not resumen:
        print("No hay datos de conteo para exportar")
        return None

    df_resumen = pd.DataFrame(list(resumen.items()), columns=["Metrica", "Valor"])

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        df_resumen.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        df_resumen.to_excel(archivo_salida, index=False)
    else:
        print("Formato no soportado. Elige 'csv' o 'xlsx'.")
        return None

    print(f"Resumen de conteos exportado exitosamente como: {archivo_salida}")
    return archivo_salida


def exportar_estadisticas_autoria(df: pd.DataFrame, nombre_archivo: str = "estadisticas_autoria",
                                  formato: str = "csv") -> str | None:
    if df is None or 'Authors' not in df.columns:
        print("No hay datos de autoría para exportar")
        return None

    df_autoria = df[['Title', 'Authors', 'Year']].copy()

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        df_autoria.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        df_autoria.to_excel(archivo_salida, index=False)
    else:
        print("Formato no soportado. Elige 'csv' o 'xlsx'.")
        return None

    print(f"Estadísticas de autoría exportadas exitosamente como: {archivo_salida}")
    return archivo_salida


def exportar_promedios_citas(df: pd.DataFrame, nombre_archivo: str = "promedios_citas") -> str | None:
    if df is None or 'Promedio_Citas_Anual' not in df.columns:
        print("No hay promedios de citas para exportar, calcula primero el impacto")
        return None

    df_exportar = df[['Title', 'Year', 'Promedio_Citas_Anual']].copy()
    df_exportar = df_exportar.sort_values(by='Promedio_Citas_Anual', ascending=False)

    archivo_salida = f"{nombre_archivo}.xlsx"
    df_exportar.to_excel(archivo_salida, index=False)

    print(f"Promedios de citas exportados exitosamente como: {archivo_salida}")
    return archivo_salida


# --- Exportaciones Múltiples y Formato APA (Historias 51, 52) ---

def exportar_secciones(df: pd.DataFrame, resumen: dict, secciones: list,
                       nombre_archivo: str = "reporte_personalizado") -> str | None:
    if not secciones:
        print("No se eligió ninguna sección para exportar")
        return None

    archivo_salida = f"{nombre_archivo}.xlsx"

    with pd.ExcelWriter(archivo_salida, engine='openpyxl') as writer:
        if 'conteos' in secciones and resumen:
            df_conteos = pd.DataFrame(list(resumen.items()), columns=["Metrica", "Valor"])
            df_conteos.to_excel(writer, sheet_name='Conteos', index=False)

        if 'autoria' in secciones and df is not None:
            df[['Title', 'Authors', 'Year']].to_excel(writer, sheet_name='Autoria', index=False)

        if 'impacto' in secciones and df is not None and 'Promedio_Citas_Anual' in df.columns:
            df[['Title', 'Year', 'Promedio_Citas_Anual']].to_excel(writer, sheet_name='Impacto', index=False)

    print(f"Reporte personalizado exportado exitosamente como: {archivo_salida}")
    return archivo_salida


def exportar_referencias_apa(df: pd.DataFrame, nombre_archivo: str = "referencias_apa",
                             formato: str = "csv") -> str | None:
    if df is None or not {'Authors', 'Title', 'Year'}.issubset(df.columns):
        print("Faltan columnas necesarias para generar referencias APA")
        return None

    df_apa = df.copy()
    df_apa['Referencia_APA'] = (
            df_apa['Authors'].fillna('Sin autor') + '. (' +
            df_apa['Year'].astype(str) + '). ' +
            df_apa['Title'].fillna('Sin título') + '.'
    )

    if 'DOI' in df_apa.columns:
        df_apa['Referencia_APA'] += df_apa['DOI'].apply(
            lambda doi: f' https://doi.org/{doi}' if pd.notna(doi) else ''
        )

    df_exportar = df_apa[['Referencia_APA']].copy()

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        df_exportar.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        df_exportar.to_excel(archivo_salida, index=False)
    else:
        print("Formato no soportado. Elige 'csv' o 'xlsx'.")
        return None

    print(f"Referencias APA exportadas exitosamente como: {archivo_salida}")
    return archivo_salida


# --- Exportación General a Excel/Word (Historias 46, 47, 49, 50, 53, 54) ---

def exportar_resultados(df: pd.DataFrame, nom_metrica: str, formato: str) -> bool:
    if df is None: return False
    if df.empty: return False

    nom = nom_metrica.replace(' ', '_')
    archivos_guardados = []

    if formato == 'excel' or formato == 'ambos':
        nom_archivo_excel = f"Analisis_{nom}.xlsx"
        df.to_excel(nom_archivo_excel, sheet_name=nom_metrica, index=False)
        archivos_guardados.append(nom_archivo_excel)

    if formato == 'word' or formato == 'ambos':
        doc = Document()
        doc.add_heading(f"Resultados del análisis: {nom_metrica}", level=1)

        tabla_de_word = doc.add_table(rows=1, cols=df.shape[1])
        tabla_de_word.style = 'Table Grid'

        for i, col in enumerate(df.columns):
            tabla_de_word.rows[0].cells[i].text = str(col)

        for _, fila in df.iterrows():
            celdas = tabla_de_word.add_row().cells
            for i, valor in enumerate(fila):
                celdas[i].text = str(valor)

        nom_archivo_word = f"Resultados del análisis_{nom}.docx"
        doc.save(nom_archivo_word)
        archivos_guardados.append(nom_archivo_word)

    print(f"Análisis exportados en el archivo:")
    for archivo in archivos_guardados:
        print(archivo)

    return True


# --- Exportar Top 10 Autores (Historia 36) ---

def exportar_top_10_autores(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty or 'Authors' not in df.columns or 'Title' not in df.columns:
        return pd.DataFrame()

    datos = []
    for _, fila in df.dropna(subset=['Authors', 'Title']).iterrows():
        autores = str(fila['Authors']).split(';')
        titulo = str(fila['Title']).strip()

        for autor in autores:
            autor = autor.strip()
            if autor:
                datos.append({'Autor': autor, 'Título': titulo})

    df_autores = pd.DataFrame(datos)

    df_agrupado = df_autores.groupby('Autor').agg(
        numero_articulos=('Título', 'count'),
        titulos_unidos=('Título', lambda x: " | ".join(x.unique()))
    ).reset_index()

    df_top = df_agrupado.sort_values(by='numero_articulos', ascending=False).head(10)
    df_top = df_top.rename(columns={
        'numero_articulos': 'Número de artículos',
        'titulos_unidos': 'Títulos'
    })

    archivo = "Los 10 autores mas destacados.csv"
    df_top.to_csv(archivo, index=False, encoding='utf-8-sig')

    return df_top

# HU 23: Exportar lista de paises
def exportar_frecuencia_paises(df_paises: pd.DataFrame, nombre_archivo: str = "frecuencia_paises",
                               formato: str = "csv") -> str | None:
    if df_paises is None or df_paises.empty:
        print("No hay datos para exportar.")
        return None

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        df_paises.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        df_paises.to_excel(archivo_salida, index=False)
    else:
        print("Formato no soportado.")
        return None

    print(f"Lista de países exportada como: {archivo_salida}")
    return archivo_salida

# HU 47: Exportar Top 10 trabajos
def exportar_top_10_trabajos(df_top_10: pd.DataFrame, nombre_archivo: str = "top_10_trabajos",
                             formato: str = "csv") -> str | None:
    if df_top_10 is None or df_top_10.empty:
        print("No hay trabajos para exportar.")
        return None

    if formato.lower() == "csv":
        archivo_salida = f"{nombre_archivo}.csv"
        df_top_10.to_csv(archivo_salida, index=False, encoding='utf-8-sig')
    elif formato.lower() == "xlsx":
        archivo_salida = f"{nombre_archivo}.xlsx"
        df_top_10.to_excel(archivo_salida, index=False)
    elif formato.lower() == "word":
        archivo_salida = f"{nombre_archivo}.docx"
        doc = Document()
        doc.add_heading("Top 10 Trabajos Más Relevantes", level=1)

        tabla = doc.add_table(rows=1, cols=len(df_top_10.columns))
        tabla.style = 'Table Grid'
        for i, col in enumerate(df_top_10.columns):
            tabla.rows[0].cells[i].text = str(col)

        for _, fila in df_top_10.iterrows():
            celdas = tabla.add_row().cells
            for i, valor in enumerate(fila):
                celdas[i].text = str(valor)
        doc.save(archivo_salida)
    else:
        print("Formato no soportado.")
        return None

    print(f"Top 10 trabajos exportado como: {archivo_salida}")
    return archivo_salida