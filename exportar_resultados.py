# HU 49,50,53,54 Exportacion de los análisis

from docx import Document

def exportar_resultados(df, nom_metrica, formato):

    if df is None: return False
    if df.empty: return False

    # Se cambian los espacios por guiones bajo
    nom = nom_metrica.replace(' ', '_')
    
    archivos_guardados = []

    if formato == 'excel' or formato == 'ambos':
        # Exportamos a excel
        nom_archivo_excel = f"Analisis_{nom}.xlsx"
        # index como falso para evitar columnas basura
        df.to_excel(nom_archivo_excel,sheet_name = nom_metrica,index=False)
        archivos_guardados.append(nom_archivo_excel)
        
    if formato == 'word' or formato == 'ambos':
        # Exportamos a word
        doc = Document()
        
        # Agregamos el título 
        doc.add_heading(f"Resultados del análisis:{nom_metrica}",level=1)
        
        # Hacemos una tabla con bordes
        tabla_de_word = doc.add_table(rows=1,cols=df.shape[1])
        tabla_de_word.style = 'Table Grid'
        
        # Se ponen los subtitulos
        for i, col in enumerate(df.columns):
            tabla_de_word.rows[0].cells[i].text = str(col)
            
        # Ahora los datos fila por fila
        for _, fila in df.iterrows():
            celdas = tabla_de_word.add_row().cells
            for i, valor in enumerate(fila):
                celdas[i].text = str(valor)
                
        # Guardan los archivos
        nom_archivo_word = f"Resultados del análisis_{nom}.docx"
        doc.save(nom_archivo_word)
        archivos_guardados.append(nom_archivo_word)
        
    print(f"Análisis exportados en el archivo:")
    for archivo in archivos_guardados:
        print(archivo)
        
    return True
